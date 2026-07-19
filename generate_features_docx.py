#!/usr/bin/env python3
"""Generate a .docx feature-summary for the Video Automation Tool."""

from docx import Document
from docx.shared import Pt, Inches

doc = Document()

# ── styles ──────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Segoe UI'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)

for level in range(1, 4):
    s = doc.styles[f'Heading {level}']
    s.font.name = 'Segoe UI'

# ── HERO ────────────────────────────────────────────────────────────
doc.add_heading('Video Automation Studio', level=1)
doc.add_heading('Almost Impossible. You Will Never Need a Video Editor Again.', level=2)

p = doc.add_paragraph()
p.add_run(
    'One click — quotes, transitions, voiceovers, captions, effects, '
    'AI lip-sync, cleaning. Two years of relentless iteration. '
    'From raw footage to platform-ready content in minutes, not hours.'
).italic = True

doc.add_paragraph()  # spacer

# ── HELPERS ─────────────────────────────────────────────────────────
def add_section(title, items):
    doc.add_heading(title, level=2)
    for bold, rest in items:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(f'{bold}  ')
        run_b.bold = True
        p.add_run(rest)

# ── CONTENT ─────────────────────────────────────────────────────────

add_section('Core Pipeline — One Click, Full Video', [
    ('Auto Quote Engine',
     'Transcript or Excel → timed text overlays with per-niche prompts '
     '(movies, courtroom, heartwarming, voiceover).'),
    ('23+ Transitions',
     'Fade, Zoom, Blur, Slide, Wipe, Glitch, Bounce, Mask Reveal, '
     'Radial Wipe, Split Wipe, Luma Wipe, Cinematic Bars, Color Dissolve '
     '— each with 15+ SFX sounds.'),
    ('Smart CTA System',
     'Auto call-to-action with spoken/shown label stripping, multi-line '
     'wrapping, freeze-frame tail extension.'),
    ('Scene-by-Scene Presets',
     'Per-niche script modes (Rewrite / Write Story) with dual prompt-store '
     'architecture.'),
    ('Multi-Platform Output',
     'Instagram Reels, TikTok, YouTube Shorts, YouTube, Facebook — '
     'custom resolution & crop modes.'),
])

add_section('Artificial Intelligence', [
    ('4 TTS Engines',
     'Microsoft Edge Cloud (41+ voices, 17 locales), Kokoro ONNX (42 voices), '
     'Qwen3-TTS (emotion control + voice cloning), NeuTTS.'),
    ('Gemini Case Commentary',
     'Watches courtroom/movie video → structured Summary + Montage Clips + '
     'Commentary Spots → auto-builds the montage.'),
    ('Whisper Speech-to-Text',
     'Word-level timing for TTS, dialogue captions from original audio, '
     'overlap suppression with voiceover.'),
    ('Wav2Lip AI Avatar',
     'Lip-sync any face image to speech, overlay or standalone mode, '
     'position & size controls.'),
    ('LLM Rerank',
     'Optional Gemini / OpenAI / Anthropic reranking over heuristic '
     'scene-cut scoring.'),
])

add_section('Captions — Three Independent Stacks', [
    ('Simple Captions',
     'Style presets, font/color/size, background, stroke, position, '
     'sync offset, emoji presets.'),
    ('Highlight Captions (CapCut Style)',
     'Word-by-word colour-highlight animation, independent font & animation type.'),
    ('Dialogue Captions (Whisper)',
     'Auto-transcribe original audio, third caption layer, smart word-level '
     'overlap suppression with TTS/commentary.'),
    ('Live Preview',
     'Phone-aspect canvas with toggle layers, scrubber, crosshair alignment guide.'),
])

add_section('Audio Engineering', [
    ('BGM System',
     'File or folder, volume, auto-loop, intelligent ducking that drops '
     'during speech.'),
    ('Voiceover Ducking',
     'Per-frame amplitude check — original audio drops to 15% during '
     'TTS segments.'),
    ('Voice Effects',
     'Deep / High / Robot / Echo / Whisper / Radio / Chipmunk + SSML '
     'speaking styles.'),
    ('Audio Normalization',
     'Consistent loudness with configurable dB target.'),
    ('Auto Silence Removal',
     'Threshold, min-duration, transition smoothing.'),
])

add_section('Visual Effects', [
    ('Particle Systems',
     'Glitter, Stars, Hearts, Confetti — all with animated motion paths.'),
    ('Light Effects',
     'Light Leaks (warm/cold/pink/purple/rainbow), Lens Flare, Film Burn '
     '— configurable intensity & repeat interval.'),
    ('Blur Engine',
     'Region blur, custom rectangular blur regions, feather edge per-pixel control.'),
    ('Alight Motion Look Builder',
     '14+ AM-style presets: hue, saturation, gamma, shadows, VHS, scanlines, '
     'RGB split, bloom, colorama, grain, vignette, film burn & more.'),
    ('20+ Overlay Elements',
     'Progress bar, watermark, vignette, dim, grain, drop shadow, glitch, '
     'chromatic aberration, text glow, neon glow, gradient, border, crosshair, '
     'spotlight, dual-video.'),
])

add_section('Performance — GPU Renderer', [
    ('OpenCV + NVENC Pipeline',
     'Auto-detects NVENC (RTX 2060+) vs CPU fallback. 30 fps 1080p at '
     '50–100 fps vs 1–2 fps MoviePy fallback.'),
    ('Phase 2 Optimizations',
     'AV1 seek optimisation, static overlay pre-merge, density-adaptive blend, '
     'vectorised feather cache. 3-min render: ~60 min → ~7 min.'),
])

add_section('Post-Processing & Cleanup', [
    ('Quick Wipe',
     'Instant platform-metadata strip, no re-encode.'),
    ('Full Spoof',
     'Deep re-encode with mirror, smart zoom, unique colour profile, '
     'pitch-shifted audio, borders.'),
    ('Cleanup Effects',
     'Speed/rotation, letterbox, colour wash, watermark, RGB glitch, '
     'Ken Burns zoom, reverb, transitions re-apply.'),
    ('YouTube Uploader',
     'Upload processed videos directly from the app.'),
])

add_section('Specialised Modules', [
    ('Our Script Tab',
     'Channel-aware single-video pipeline, per-channel presets, batch + skip-done.'),
    ('Xiaohongshu / RedNote Scraper',
     'Selenium + CDP API interception for hidden note IDs, cookies.txt auth, '
     'cursor pagination.'),
    ('Instagram Auth',
     'Unified cookies.txt for instaloader + yt-dlp, auto re-login on session expiry.'),
    ('AI Avatar Tab',
     'Full Wav2Lip pipeline: face image, TTS audio, lip-sync, composite back '
     'onto video.'),
])

add_section('Quality of Life', [
    ('Dark-Theme GUI',
     'Scrollable canvases, 10 notebook tabs, color pickers everywhere, '
     'live preview.'),
    ('Preset Manager',
     'Save/load/delete full effect combinations as named templates.'),
    ('Platform Quick-Select',
     'One-click resolution presets with live aspect-ratio display.'),
    ('Settings Persistence',
     'overlay_settings.json + processing_paths.json auto-save/load.'),
    ('Portable Copy',
     'Standalone deployment ready on any machine.'),
])

# ── CLOSING ─────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run(
    'Four TTS engines, five AI integrations, 23 transitions, three caption stacks, '
    'a GPU render pipeline, and a cleaning engine that strips every fingerprint.'
).italic = True
p.add_run('\nAll behind one button.').bold = True

p2 = doc.add_paragraph()
p2.add_run('One Click. Ready to Post.').bold = True

# ── SAVE ───────────────────────────────────────────────────────────
out = r'd:\GitHub\ChangeGUI\Video_Automation_Studio_Features.docx'
doc.save(out)
print(f'DOCX saved to {out}')
